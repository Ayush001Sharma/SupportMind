"""
conftest.py — Shared pytest fixtures for SupportMind tests.

All external dependencies (Ollama, ChromaDB) are mocked here so that
no real inference or disk I/O happens during the test suite.
"""

import os
import uuid
from unittest.mock import MagicMock

import pytest

# ------------------------------------------------------------------ #
# Set dummy env vars BEFORE any app import so pydantic-settings
# validates cleanly without a real .env file present.
# ------------------------------------------------------------------ #
os.environ.setdefault("APP_NAME", "SupportMind-Test")
os.environ.setdefault("OLLAMA_BASE_URL", "http://localhost:11434")
os.environ.setdefault("OLLAMA_EMBEDDING_MODEL", "nomic-embed-text")
os.environ.setdefault("OLLAMA_CHAT_MODEL", "llama3.2:3b")
os.environ.setdefault("CHROMA_PERSIST_DIR", "/tmp/chromadb_test")
os.environ.setdefault("CHROMA_COLLECTION_NAME", "test_collection")


# ------------------------------------------------------------------ #
# Fixtures: Core infrastructure mocks
# ------------------------------------------------------------------ #

@pytest.fixture
def mock_embeddings():
    """Returns a fake LangChain Embeddings client that produces deterministic vectors."""
    mock = MagicMock()
    mock.embed_query.return_value = [0.1] * 384
    mock.embed_documents.return_value = [[0.1] * 384, [0.2] * 384]
    return mock


@pytest.fixture
def mock_collection():
    """Returns a fake ChromaDB Collection."""
    mock = MagicMock()
    mock.query.return_value = {
        "ids": [[]],
        "distances": [[]],
        "documents": [[]],
        "metadatas": [[]],
    }
    mock.add.return_value = None
    mock.count.return_value = 0
    return mock


@pytest.fixture
def mock_chat_model():
    """Returns a fake LangChain ChatModel that returns a deterministic answer."""
    mock = MagicMock()
    response = MagicMock()
    response.content = "The return policy allows returns within 30 days."
    mock.invoke.return_value = response
    return mock


@pytest.fixture
def settings():
    """Returns the real Settings object (driven by os.environ set above)."""
    from app.core.config import get_settings
    get_settings.cache_clear()
    return get_settings()


# ------------------------------------------------------------------ #
# Fixtures: Document helpers
# ------------------------------------------------------------------ #

@pytest.fixture
def sample_txt_bytes():
    text = "Customer support is available Monday to Friday from 9am to 5pm."
    return text.encode("utf-8")


@pytest.fixture
def minimal_pdf_bytes():
    """Returns the bytes of the smallest valid PDF containing one word."""
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>\nstream\nBT /F1 12 Tf 100 700 Td (Hello PDF) Tj ET\nendstream\nendobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000360 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n441\n%%EOF"
    )
    return pdf


@pytest.fixture
def minimal_docx_bytes():
    """Returns bytes of a minimal valid DOCX file."""
    import io
    from docx import Document as DocxDocument
    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("This is a DOCX test document.")
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ------------------------------------------------------------------ #
# Fixtures: ProcessedDocument factory
# ------------------------------------------------------------------ #

@pytest.fixture
def processed_document():
    """Returns a complete ProcessedDocument for use in chunking/indexing tests."""
    from app.schemas.processing import PageMetadata, ProcessedDocument, ProcessedPage

    page_meta = PageMetadata(
        filename="test.txt",
        page_number=1,
        document_type="txt",
        processing_timestamp="2026-07-06T00:00:00Z",
        processing_duration_ms=5.0,
    )
    page = ProcessedPage(
        page_number=1,
        text="Customer support hours are 9am to 5pm Monday through Friday. "
             "We handle refunds, product issues, shipping queries, and account changes.",
        character_count=100,
        metadata=page_meta,
    )
    return ProcessedDocument(
        document_id=str(uuid.uuid4()),
        filename="test.txt",
        document_type="txt",
        upload_timestamp="2026-07-06T00:00:00Z",
        total_pages=1,
        pages=[page],
    )


# ------------------------------------------------------------------ #
# Fixtures: RetrievalResult factory
# ------------------------------------------------------------------ #

@pytest.fixture
def retrieval_result_with_chunks():
    """Returns a RetrievalResult containing one high-confidence chunk."""
    from app.schemas.retrieval import RetrievalResult, RetrievedChunk

    chunk = RetrievedChunk(
        chunk_id="abc123",
        text="Returns are accepted within 30 days of purchase with a receipt.",
        similarity_score=0.87,
        filename="returns_policy.pdf",
        page_number=3,
        metadata={"filename": "returns_policy.pdf", "page_number": 3},
    )
    return RetrievalResult(
        query="What is the return policy?",
        retrieved_chunks=[chunk],
        total_chunks=1,
        max_similarity_score=0.87,
        retrieval_duration_ms=45.0,
    )


@pytest.fixture
def empty_retrieval_result():
    """Returns a RetrievalResult with zero chunks (triggers fallback)."""
    from app.schemas.retrieval import RetrievalResult

    return RetrievalResult(
        query="What is the weather?",
        retrieved_chunks=[],
        total_chunks=0,
        max_similarity_score=0.0,
        retrieval_duration_ms=12.0,
    )


# ------------------------------------------------------------------ #
# Fixtures: FastAPI TestClient with mocked dependencies
# ------------------------------------------------------------------ #

@pytest.fixture
def test_client(mock_embeddings, mock_collection, mock_chat_model, settings):
    """
    TestClient with all external I/O mocked.
    Overrides FastAPI DI so no Ollama or ChromaDB connections are made.
    """
    from fastapi.testclient import TestClient
    from app.main import create_app
    from app.api.dependencies import (
        get_ollama_embeddings,
        get_ollama_chat_model,
        get_chroma_collection,
    )
    from app.core.config import get_settings

    app = create_app()

    # Override all external service dependencies
    app.dependency_overrides[get_settings] = lambda: settings
    app.dependency_overrides[get_ollama_embeddings] = lambda: mock_embeddings
    app.dependency_overrides[get_chroma_collection] = lambda: mock_collection
    app.dependency_overrides[get_ollama_chat_model] = lambda: mock_chat_model

    with TestClient(app, raise_server_exceptions=True) as client:
        yield client

    app.dependency_overrides.clear()
